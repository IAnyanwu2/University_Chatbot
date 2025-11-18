import os
import logging
import json
import pickle
import time
import random
from typing import List, Dict, Any
from pathlib import Path
from datetime import datetime, timedelta
import requests
from simple_langchain import Document

# Optional imports for additional file types
try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document ingestion from various sources with persistent caching"""
    
    def __init__(self, data_dir: str = "./data"):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(exist_ok=True)
        self.cache_file = self.data_dir / "scraped_documents.pkl"
        self.cache_metadata_file = self.data_dir / "cache_metadata.json"
        self.refresh_log_file = self.data_dir / "refresh_log.json"
    
    def load_real_gsu_data(self) -> List[Document]:
        """Load GSU CS data from persistent cache (always fast, never blocks users)"""
        
        # Always try to load from cache first - never block users for scraping
        if self.cache_file.exists():
            logger.info("Loading GSU data from persistent cache...")
            try:
                with open(self.cache_file, 'rb') as f:
                    cached_documents = pickle.load(f)
                logger.info(f"Successfully loaded {len(cached_documents)} cached documents")
                return cached_documents
            except Exception as e:
                logger.warning(f"Failed to load cache: {e}. Falling back to fresh scraping.")
        
        # Only scrape if no cache exists at all (first time setup)
        logger.info("No cache found. Performing initial data scraping...")
        documents = self._scrape_fresh_data()
        self._save_to_cache(documents)
        return documents
    
    def get_cache_info(self) -> Dict[str, Any]:
        """Get information about the current cache"""
        if not self.cache_metadata_file.exists():
            return {"status": "no_cache", "document_count": 0}
        
        try:
            with open(self.cache_metadata_file, 'r') as f:
                metadata = json.load(f)
            
            cache_time = datetime.fromisoformat(metadata['cached_at'])
            age_hours = (datetime.now() - cache_time).total_seconds() / 3600
            
            return {
                "status": "available",
                "document_count": metadata.get('document_count', 0),
                "cached_at": metadata['cached_at'],
                "age_hours": round(age_hours, 1),
                "cache_size_mb": round(self.cache_file.stat().st_size / (1024*1024), 2) if self.cache_file.exists() else 0
            }
        except Exception as e:
            logger.error(f"Error reading cache metadata: {e}")
            return {"status": "error", "error": str(e)}
    
    def scheduled_background_refresh(self) -> Dict[str, Any]:
        """Background refresh method - call this from scheduler/cron job"""
        refresh_start = datetime.now()
        logger.info("Starting scheduled background cache refresh...")
        
        try:
            # Scrape fresh data
            documents = self._scrape_fresh_data()
            
            # Save to cache
            self._save_to_cache(documents)
            
            # Log successful refresh
            refresh_end = datetime.now()
            duration = (refresh_end - refresh_start).total_seconds()
            
            refresh_log = {
                "timestamp": refresh_end.isoformat(),
                "status": "success",
                "duration_seconds": round(duration, 1),
                "documents_updated": len(documents),
                "urls_scraped": 62  # Current URL count
            }
            
            # Save refresh log
            self._save_refresh_log(refresh_log)
            
            logger.info(f"Background refresh completed successfully in {duration:.1f} seconds")
            return refresh_log
            
        except Exception as e:
            # Log failed refresh
            refresh_end = datetime.now()
            duration = (refresh_end - refresh_start).total_seconds()
            
            refresh_log = {
                "timestamp": refresh_end.isoformat(),
                "status": "failed",
                "duration_seconds": round(duration, 1),
                "error": str(e),
                "documents_updated": 0
            }
            
            self._save_refresh_log(refresh_log)
            
            logger.error(f"Background refresh failed after {duration:.1f} seconds: {e}")
            return refresh_log
    
    def _save_refresh_log(self, refresh_log: Dict[str, Any]) -> None:
        """Save refresh log for monitoring"""
        try:
            # Load existing logs
            logs = []
            if self.refresh_log_file.exists():
                with open(self.refresh_log_file, 'r') as f:
                    logs = json.load(f)
            
            # Add new log
            logs.append(refresh_log)
            
            # Keep only last 30 refresh attempts
            logs = logs[-30:]
            
            # Save back
            with open(self.refresh_log_file, 'w') as f:
                json.dump(logs, f, indent=2)
                
        except Exception as e:
            logger.warning(f"Failed to save refresh log: {e}")
    
    def get_refresh_history(self) -> List[Dict[str, Any]]:
        """Get recent refresh history for monitoring"""
        try:
            if self.refresh_log_file.exists():
                with open(self.refresh_log_file, 'r') as f:
                    return json.load(f)
            return []
        except Exception as e:
            logger.error(f"Error reading refresh history: {e}")
            return []
    
    def _scrape_fresh_data(self) -> List[Document]:
        """Scrape fresh data from GSU websites"""
        documents = []
        
        # GSU CS graduate program URLs to scrape (comprehensive coverage)
        gsu_urls = [
            # Core program pages
            "https://csds.gsu.edu/graduate/",
            "https://csds.gsu.edu/graduate-faqs/", 
            "https://cas.gsu.edu/program/computer-science-phd/",
            "https://cas.gsu.edu/program/data-science-and-analytics-ms-big-data-and-machine-learning/",
            "https://csds.gsu.edu/research/",
            
            # Faculty directory and research infrastructure
            "https://csds.gsu.edu/directory/",
            "https://csds.gsu.edu/research-groups-labs/",
            "https://csds.gsu.edu/undergraduate-research/",
            
            # All faculty profiles for complete coverage
            "https://csds.gsu.edu/profile/berkay-aydin/",        # Database Systems
            "https://csds.gsu.edu/profile/anu-bourgeois/",       # Algorithms, Networks
            "https://csds.gsu.edu/profile/yanqing-zhang/",       # AI, Machine Learning
            "https://csds.gsu.edu/profile/ashwin-ashok/",        # Systems, IoT, Networks
            "https://csds.gsu.edu/profile/ying-zhu/",            # Computer Vision, Graphics
            "https://csds.gsu.edu/profile/wei-li/",              # Software Engineering
            "https://csds.gsu.edu/profile/yubao-wu/",            # Data Mining, Security
            "https://csds.gsu.edu/profile/jingyu-liu/",          # High Performance Computing
            "https://csds.gsu.edu/profile/raj-sunderraman/",     # AI, Logic Programming
            "https://csds.gsu.edu/profile/zhipeng-cai/",         # Algorithms, Networks
            "https://csds.gsu.edu/profile/xiaolin-hu/",          # AI, Neural Networks
            "https://csds.gsu.edu/profile/xiaojun-cao/",         # Graduate Director
            "https://csds.gsu.edu/profile/murray-patterson/",    # Bioinformatics
            "https://csds.gsu.edu/profile/robyn-miller/",        # Software Engineering
            "https://csds.gsu.edu/profile/michael-weeks/",       # Image Processing
            
            # Additional faculty profiles discovered
            "https://csds.gsu.edu/profile/sergey-morozov/",      # Systems, Security
            "https://csds.gsu.edu/profile/xiaojuan-ma/",         # HCI, Interactive Systems
            "https://csds.gsu.edu/profile/yi-pan/",              # Algorithms, Bioinformatics
            "https://csds.gsu.edu/profile/saeid-belkasim/",      # Image Processing, AI
            "https://csds.gsu.edu/profile/alex-zelikovsky/",     # Algorithms, Bioinformatics
            
            # Faculty research and publication pages
            "https://csds.gsu.edu/profile/berkay-aydin/#profile-publications-tab",
            "https://csds.gsu.edu/profile/anu-bourgeois/#profile-publications-tab",
            "https://csds.gsu.edu/profile/yanqing-zhang/#profile-publications-tab",
            "https://csds.gsu.edu/profile/ashwin-ashok/#profile-publications-tab",
            "https://csds.gsu.edu/profile/ying-zhu/#profile-publications-tab",
            "https://csds.gsu.edu/profile/wei-li/#profile-publications-tab",
            "https://csds.gsu.edu/profile/yubao-wu/#profile-publications-tab",
            "https://csds.gsu.edu/profile/jingyu-liu/#profile-publications-tab",
            "https://csds.gsu.edu/profile/raj-sunderraman/#profile-publications-tab",
            "https://csds.gsu.edu/profile/zhipeng-cai/#profile-publications-tab",
            
            # Program administration and details
            "https://csds.gsu.edu/graduate/advisement/",
            "https://csds.gsu.edu/course-descriptions-schedule/",
            "https://csds.gsu.edu/graduate/martin-d-fraser-graduate-student-conference-travel-awards/",
            "https://csds.gsu.edu/undergraduate/advisement/",
            "https://csds.gsu.edu/undergraduate/co-op/",
            "https://csds.gsu.edu/undergraduate-faqs/",
            
            # Research areas and specializations
            "https://csds.gsu.edu/research/#hardware-and-software-systems",
            "https://csds.gsu.edu/research/#theory-of-computation",
            "https://csds.gsu.edu/research/#artificial-intelligence",
            "https://csds.gsu.edu/research/#software-engineering",
            "https://csds.gsu.edu/research/#database-and-data-mining",
            
            # Department information and contact
            "https://csds.gsu.edu/about/",
            "https://csds.gsu.edu/faculty-staff-forms/",
            "https://csds.gsu.edu/news/",
            "https://csds.gsu.edu/events/",
            
            # Course and curriculum information
            "https://csds.gsu.edu/two-year-course-schedule/",
            "https://csds.gsu.edu/cs-prerequisite-chart/",
            "https://csds.gsu.edu/undergraduate/gateway-to-computer-science/",
            
            # Additional program pages from CAS
            "https://cas.gsu.edu/program/computer-science-ms/",
            "https://cas.gsu.edu/program/computer-science-bs/",
            
            # Student resources and support
            "https://csds.gsu.edu/student-organizations/",
            "https://csds.gsu.edu/scholarships-awards/",
            "https://csds.gsu.edu/internships-careers/"
        ]
        
        logger.info(f"Scraping {len(gsu_urls)} verified GSU websites for fresh data...")
        
        for i, url in enumerate(gsu_urls):
            try:
                # Add random delay between requests to avoid rate limiting
                if i > 0:
                    delay = random.uniform(2, 5)
                    logger.info(f"Waiting {delay:.1f} seconds before next request...")
                    time.sleep(delay)
                
                logger.info(f"Scraping: {url}")
                doc = self.scrape_webpage(url)
                
                # Only add if scraping was successful
                if doc.page_content and "Error scraping" not in doc.page_content:
                    documents.append(doc)
                    logger.info(f"Successfully scraped {url} - {len(doc.page_content)} characters")
                else:
                    logger.warning(f"Failed to scrape meaningful content from {url}")
                    
            except Exception as e:
                logger.error(f"Error scraping {url}: {e}")
                continue
        
        # Also load any PDF documents from the data directory
        try:
            pdf_docs = self.load_documents_from_directory(str(self.data_dir))
            if pdf_docs:
                documents.extend(pdf_docs)
                logger.info(f"Added {len(pdf_docs)} documents from {self.data_dir}")
        except Exception as e:
            logger.warning(f"Could not load documents from {self.data_dir}: {e}")
        
        if not documents:
            logger.error("No documents were successfully loaded! Check network connection and URLs")
            # Fallback to prevent complete failure
            return [Document(
                page_content="No data available. Please contact the CS department directly at cs-grad@gsu.edu or (404) 413-5820.",
                metadata={"source": "fallback", "category": "error", "last_updated": "2025-11-02"}
            )]
        
        logger.info(f"Successfully scraped {len(documents)} fresh GSU documents")
        return documents
    
    def _save_to_cache(self, documents: List[Document]) -> None:
        """Save documents to cache"""
        try:
            # Save documents
            with open(self.cache_file, 'wb') as f:
                pickle.dump(documents, f)
            
            # Save metadata
            metadata = {
                'cached_at': datetime.now().isoformat(),
                'document_count': len(documents),
                'refresh_type': 'manual'  # vs 'scheduled'
            }
            with open(self.cache_metadata_file, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            logger.info(f"Cached {len(documents)} documents persistently.")
        except Exception as e:
            logger.warning(f"Failed to save cache: {e}")
    
    def force_refresh_cache(self) -> List[Document]:
        """Force refresh the cache by scraping fresh data (for manual/admin use)"""
        logger.info("Force refreshing cache...")
        documents = self._scrape_fresh_data()
        self._save_to_cache(documents)
        return documents
    
    def load_text_file(self, file_path: str) -> Document:
        """Load a plain text file"""
        path = Path(file_path)
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return Document(
            page_content=content,
            metadata={"source": path.stem, "file_type": "text", "file_path": str(path)}
        )
    
    def load_pdf_file(self, file_path: str) -> Document:
        """Load a PDF file"""
        if not HAS_PDF:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return Document(
                page_content="PDF processing not available - PyPDF2 not installed",
                metadata={"source": Path(file_path).stem, "file_type": "pdf", "error": "missing_dependency"}
            )
        
        path = Path(file_path)
        content = ""
        
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                content += page.extract_text() + "\\n"
        
        return Document(
            page_content=content.strip(),
            metadata={"source": path.stem, "file_type": "pdf", "file_path": str(path)}
        )
    
    def load_docx_file(self, file_path: str) -> Document:
        """Load a Word document"""
        if not HAS_DOCX:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return Document(
                page_content="DOCX processing not available - python-docx not installed",
                metadata={"source": Path(file_path).stem, "file_type": "docx", "error": "missing_dependency"}
            )
        
        path = Path(file_path)
        doc = DocxDocument(path)
        content = "\\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return Document(
            page_content=content,
            metadata={"source": path.stem, "file_type": "docx", "file_path": str(path)}
        )
    
    def scrape_webpage(self, url: str, selector: str = None) -> Document:
        """Scrape content from a webpage with enhanced anti-bot measures"""
        if not HAS_BS4:
            logger.error("BeautifulSoup4 not installed. Install with: pip install beautifulsoup4")
            return Document(
                page_content="Web scraping not available - beautifulsoup4 not installed",
                metadata={"source": url, "file_type": "webpage", "error": "missing_dependency"}
            )
        
        try:
            # Enhanced headers to mimic real browser
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1'
            }
            
            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Extract text based on selector or get all text
            if selector:
                elements = soup.select(selector)
                content = "\n".join([elem.get_text().strip() for elem in elements])
            else:
                content = soup.get_text()
            
            # Clean up text aggressively but keep it simple
            content = soup.get_text()
            lines = (line.strip() for line in content.splitlines())
            content = "\n".join(line for line in lines if line)
            
            # Basic cleanup: normalize whitespace
            import re
            content = re.sub(r'\s+', ' ', content)  # Normalize whitespace
            content = content.strip()
            
            # Verify we got readable content
            if len(content) < 200:
                logger.warning(f"Scraped content too short for {url}")
                return Document(
                    page_content=f"Failed to extract sufficient content from {url}",
                    metadata={"source": url, "file_type": "webpage", "error": "content_too_short"}
                )
            
            return Document(
                page_content=content,
                metadata={"source": url, "file_type": "webpage", "scraped_at": "2024-10-21"}
            )
            
        except Exception as e:
            logger.error(f"Error scraping {url}: {e}")
            return Document(
                page_content=f"Error scraping content from {url}",
                metadata={"source": url, "file_type": "webpage", "error": str(e)}
            )
    
    def load_documents_from_directory(self, directory: str) -> List[Document]:
        """Load all supported documents from a directory"""
        directory = Path(directory)
        documents = []
        
        if not directory.exists():
            logger.warning(f"Directory {directory} does not exist")
            return documents
        
        for file_path in directory.rglob("*"):
            if file_path.is_file():
                try:
                    if file_path.suffix.lower() == '.txt':
                        documents.append(self.load_text_file(str(file_path)))
                    elif file_path.suffix.lower() == '.pdf':
                        documents.append(self.load_pdf_file(str(file_path)))
                    elif file_path.suffix.lower() in ['.docx', '.doc']:
                        documents.append(self.load_docx_file(str(file_path)))
                    else:
                        logger.info(f"Skipping unsupported file type: {file_path}")
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {e}")
        
        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents